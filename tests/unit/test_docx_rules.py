"""Regression tests for Word/DOCX detection precision."""

from __future__ import annotations

import io

from docx import Document  # type: ignore[import-untyped]
from docx.oxml import OxmlElement, parse_xml  # type: ignore[import-untyped]
from docx.oxml.ns import qn as docx_qn
from PIL import Image  # type: ignore[import-untyped]

from a11yfix.ooxml.docx_reader import open_docx
from a11yfix.ooxml.image_extract import extract_image_for_finding
from a11yfix.rules.alt_text import AltTextRule
from a11yfix.rules.document_language import DocumentLanguageRule
from a11yfix.rules.heading_structure import HeadingStructureRule
from a11yfix.rules.link_text import LinkTextRule
from a11yfix.rules.list_semantics import ListSemanticsRule
from a11yfix.rules.merged_cells import MergedCellsRule
from a11yfix.rules.table_headers import TableHeaderRule


def test_docx_fake_list_inside_table_cell_uses_cell_path(tmp_path):
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].text = "• typed bullet in a layout cell"
    path = tmp_path / "nested_fake_list.docx"
    doc.save(path)

    findings = list(ListSemanticsRule().detect(open_docx(path)))

    assert len(findings) == 1
    assert findings[0].officecli_path == "/body/tbl[1]/tr[1]/tc[1]/p[1]"


def test_docx_fake_heading_inside_table_cell_uses_cell_path(tmp_path):
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    paragraph = table.cell(0, 0).paragraphs[0]
    run = paragraph.add_run("Visual heading in a table")
    run.bold = True
    path = tmp_path / "nested_fake_heading.docx"
    doc.save(path)

    findings = list(HeadingStructureRule().detect(open_docx(path)))

    assert len(findings) == 1
    assert findings[0].officecli_path == "/body/tbl[1]/tr[1]/tc[1]/p[1]"


def test_docx_nested_table_does_not_emit_second_body_table_path(tmp_path):
    doc = Document()
    outer = doc.add_table(rows=2, cols=2)
    outer.cell(0, 0).add_table(rows=1, cols=1)
    path = tmp_path / "nested_table.docx"
    doc.save(path)

    findings = list(TableHeaderRule().detect(open_docx(path)))

    assert len(findings) == 1
    assert findings[0].officecli_path == "/body/tbl[1]/tr[1]"


def test_docx_decorative_docpr_image_is_not_missing_alt(tmp_path):
    img = tmp_path / "image.png"
    Image.new("RGB", (50, 50), color="blue").save(img)
    doc = Document()
    inline_shape = doc.add_picture(str(img))
    doc_pr = inline_shape._inline.docPr
    extlst = parse_xml(
        '<a:extLst xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:adec="http://schemas.microsoft.com/office/drawing/2017/decorative">'
        '<a:ext uri="{C183D7F6-B498-43B3-948B-1728B52AA6E4}">'
        '<adec:decorative val="1"/></a:ext></a:extLst>'
    )
    doc_pr.append(extlst)
    path = tmp_path / "decorative.docx"
    doc.save(path)

    findings = list(AltTextRule().detect(open_docx(path)))

    assert findings == []


def test_docx_image_inside_table_cell_is_missing_alt_with_extractable_bytes(tmp_path):
    img = tmp_path / "image.png"
    Image.new("RGB", (50, 50), color="red").save(img)
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run().add_picture(str(img))
    path = tmp_path / "table_image.docx"
    doc.save(path)

    handle = open_docx(path)
    findings = list(AltTextRule().detect(handle))

    assert len(findings) == 1
    assert findings[0].officecli_path == "/body/tbl[1]/tr[1]/tc[1]/p[1]/r[1]"
    assert extract_image_for_finding(handle, findings[0]) is not None


def test_docx_hyperlinked_image_is_missing_alt_with_hyperlink_path(tmp_path):
    img = tmp_path / "image.png"
    Image.new("RGB", (50, 50), color="yellow").save(img)
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(str(img))
    hyperlink = OxmlElement("w:hyperlink")
    paragraph._p.remove(run._r)
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)
    path = tmp_path / "hyperlinked_image.docx"
    doc.save(path)

    handle = open_docx(path)
    findings = list(AltTextRule().detect(handle))

    assert len(findings) == 1
    assert findings[0].officecli_path == "/body/p[1]/hyperlink[1]/r[1]"
    assert extract_image_for_finding(handle, findings[0]) is not None


def test_docx_generic_link_inside_table_cell_uses_cell_path(tmp_path):
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    paragraph = table.cell(0, 0).paragraphs[0]
    hyperlink = OxmlElement("w:hyperlink")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "click here"
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    path = tmp_path / "table_link.docx"
    doc.save(path)

    findings = list(LinkTextRule().detect(open_docx(path)))

    assert len(findings) == 1
    assert findings[0].officecli_path == "/body/tbl[1]/tr[1]/tc[1]/p[1]/hyperlink[1]"


def test_docx_nested_merged_table_uses_nested_table_path(tmp_path):
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    inner = outer.cell(0, 0).add_table(rows=1, cols=2)
    inner.cell(0, 0).merge(inner.cell(0, 1))
    path = tmp_path / "nested_merged.docx"
    doc.save(path)

    findings = list(MergedCellsRule().detect(open_docx(path)))

    assert [f.officecli_path for f in findings] == ["/body/tbl[1]/tr[1]/tc[1]/tbl[1]"]


def test_docx_tblheader_false_is_not_treated_as_header(tmp_path):
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(docx_qn("w:val"), "false")
    tr_pr.append(tbl_header)
    path = tmp_path / "tblheader_false.docx"
    doc.save(path)

    findings = list(TableHeaderRule().detect(open_docx(path)))

    assert len(findings) == 1
    assert findings[0].officecli_path == "/body/tbl[1]/tr[1]"


def test_docx_false_bold_first_row_is_not_visual_header(tmp_path):
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    for cell in table.rows[0].cells:
        run = cell.paragraphs[0].add_run("Not a header")
        run.bold = True
        b = run._r.rPr.find(docx_qn("w:b"))
        b.set(docx_qn("w:val"), "false")
    path = tmp_path / "false_bold.docx"
    doc.save(path)

    handle = open_docx(path)
    finding = next(iter(TableHeaderRule().detect(handle)))

    assert finding.extra["visually_header"] is False
    assert TableHeaderRule().fix_deterministic(finding, handle) is None


def test_docx_nested_table_bold_does_not_make_outer_row_visual_header(tmp_path):
    doc = Document()
    outer = doc.add_table(rows=2, cols=2)
    inner = outer.cell(0, 0).add_table(rows=1, cols=1)
    inner_run = inner.cell(0, 0).paragraphs[0].add_run("Nested bold")
    inner_run.bold = True
    path = tmp_path / "nested_bold.docx"
    doc.save(path)

    handle = open_docx(path)
    finding = next(iter(TableHeaderRule().detect(handle)))

    assert finding.officecli_path == "/body/tbl[1]/tr[1]"
    assert finding.extra["visually_header"] is False
    assert TableHeaderRule().fix_deterministic(finding, handle) is None


def test_docx_direct_hyperlink_bold_can_be_visual_header(tmp_path):
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    for cell in table.rows[0].cells:
        paragraph = cell.paragraphs[0]
        hyperlink = OxmlElement("w:hyperlink")
        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rpr.append(OxmlElement("w:b"))
        text = OxmlElement("w:t")
        text.text = "Header"
        run.append(rpr)
        run.append(text)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)
    path = tmp_path / "hyperlink_bold_header.docx"
    doc.save(path)

    handle = open_docx(path)
    finding = next(iter(TableHeaderRule().detect(handle)))

    assert finding.extra["visually_header"] is True
    assert TableHeaderRule().fix_deterministic(finding, handle) is not None


def test_docx_nested_table_shading_does_not_make_outer_row_visual_header(tmp_path):
    doc = Document()
    outer = doc.add_table(rows=2, cols=2)
    inner = outer.cell(0, 0).add_table(rows=1, cols=1)
    tc_pr = inner.cell(0, 0)._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(docx_qn("w:fill"), "CCCCCC")
    tc_pr.append(shd)
    path = tmp_path / "nested_shading.docx"
    doc.save(path)

    handle = open_docx(path)
    finding = next(iter(TableHeaderRule().detect(handle)))

    assert finding.officecli_path == "/body/tbl[1]/tr[1]"
    assert finding.extra["visually_header"] is False
    assert TableHeaderRule().fix_deterministic(finding, handle) is None


def test_docx_image_extract_uses_run_path_before_duplicate_pic_id(tmp_path):
    red = tmp_path / "red.png"
    green = tmp_path / "green.png"
    Image.new("RGB", (10, 10), color="red").save(red)
    Image.new("RGB", (10, 10), color="green").save(green)
    doc = Document()
    doc.add_paragraph().add_run().add_picture(str(red))
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run().add_picture(str(green))
    for doc_pr in doc.element.body.iter(docx_qn("wp:docPr")):
        doc_pr.set("id", "42")
    path = tmp_path / "duplicate_docpr_id.docx"
    doc.save(path)

    handle = open_docx(path)
    findings = list(AltTextRule().detect(handle))
    table_finding = next(f for f in findings if "/tbl[1]/" in f.officecli_path)
    extracted = extract_image_for_finding(handle, table_finding)

    assert extracted is not None
    image = Image.open(io.BytesIO(extracted[0]))
    assert image.getpixel((0, 0)) == (0, 128, 0)


def test_docx_styles_default_language_suppresses_language_missing(tmp_path):
    doc = Document()
    doc.add_paragraph("Hello")
    styles = doc.styles.element
    doc_defaults = styles.find(docx_qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(docx_qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)
    rpr = rpr_default.find(docx_qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    lang = rpr.find(docx_qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(docx_qn("w:val"), "en-US")
    path = tmp_path / "styles_lang.docx"
    doc.save(path)

    findings = list(DocumentLanguageRule().detect(open_docx(path)))

    assert findings == []
